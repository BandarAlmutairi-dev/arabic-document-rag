from io import BytesIO

from docx import Document
from pypdf import PdfWriter

from app.document_parser import (
    extract_pages_from_pdf,
    extract_paragraphs_from_docx,
    extract_text_from_docx,
    extract_text_from_pdf,
)


def test_extract_text_from_docx():
    document = Document()
    document.add_paragraph("مرحبا بالعالم")
    document.add_paragraph("هذا مستند عربي")

    buffer = BytesIO()
    document.save(buffer)

    text = extract_text_from_docx(buffer.getvalue())

    assert "مرحبا بالعالم" in text
    assert "هذا مستند عربي" in text


def test_extract_text_from_valid_pdf():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)

    buffer = BytesIO()
    writer.write(buffer)

    text = extract_text_from_pdf(buffer.getvalue())

    assert text == ""


def test_extract_paragraphs_from_docx_with_numbers():
    document = Document()
    document.add_paragraph("الفقرة الأولى")
    document.add_paragraph("")
    document.add_paragraph("الفقرة الثالثة")

    buffer = BytesIO()
    document.save(buffer)

    paragraphs = extract_paragraphs_from_docx(buffer.getvalue())

    assert paragraphs == [
        {
            "paragraph_number": 1,
            "text": "الفقرة الأولى",
        },
        {
            "paragraph_number": 3,
            "text": "الفقرة الثالثة",
        },
    ]


def test_extract_pages_from_blank_pdf():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)

    buffer = BytesIO()
    writer.write(buffer)

    pages = extract_pages_from_pdf(buffer.getvalue())

    assert pages == []